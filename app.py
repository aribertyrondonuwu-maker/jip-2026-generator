import io
import docx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
import streamlit as st

# ---------------------------------------------------------
# KONFIGURASI HALAMAN
# ---------------------------------------------------------
st.set_page_config(
    page_title="JIP 2026 Article Auto-Generator",
    page_icon="📄",
    layout="wide",
)

# ---------------------------------------------------------
# SIDEBAR - PILIHAN BAHASA & MODE NASKAH
# ---------------------------------------------------------
st.sidebar.header("⚙️ Settings / Pengaturan")

# Pilihan Bahasa Interface
language = st.sidebar.selectbox(
    "🌐 Select Language / Pilih Bahasa:",
    ["Bahasa Indonesia", "English"],
    index=0
)

st.sidebar.markdown("---")

# Teks UI Berdasarkan Bahasa
if language == "English":
    # English UI Texts
    ui_title = "📄 JIP 2026 Article Auto-Generator"
    ui_desc = "Fill out the form below to generate a `.docx` manuscript ready for submission according to PLATAX Scientific Journal guidelines."
    
    ui_mode_header = "Manuscript Output Mode"
    ui_mode_label = "Select Output Document Version:"
    ui_mode_opts = ["Blind Review", "Final (Full Version)"]
    ui_mode_help = "Select 'Blind Review' for initial submission, or 'Final (Full Version)' if accepted."
    
    ui_guide_title = "📌 **Submission Guidelines:**\n\n"
    ui_guide_body = (
        "• **Blind Review**: Required for **Initial Submission & Review Process**. Author info and acknowledgements are automatically hidden.\n\n"
        "• **Final (Full Version)**: Use after the manuscript is **Accepted** for publication layout."
    )
    
    tab_names = ["📌 Title & Authors", "📝 Abstract", "📄 Chapters 1–4", "🤝 Declarations & References"]
    
    lbl_title_id = "Indonesian Title (Max 20 Words)"
    lbl_title_en = "English Title (Max 20 Words)"
    lbl_running = "Running Title (Max 60 Characters)"
    lbl_author_sec = "Author Information"
    lbl_authors = "Author Names (Separate with commas)"
    lbl_affil = "Affiliation / Institution"
    lbl_email = "Corresponding Email"
    
    lbl_abs_id = "Abstract (Indonesian)"
    lbl_kw_id = "Keywords (Indonesian, separate with commas)"
    lbl_abs_en = "Abstract (English)"
    lbl_kw_en = "Keywords (English, separate with commas)"
    
    lbl_ch1 = "Chapter 1: Introduction"
    lbl_ch2 = "Chapter 2: Research Methods"
    lbl_ch3 = "Chapter 3: Results and Discussion"
    lbl_ch4 = "Chapter 4: Conclusion and Suggestions"
    
    lbl_ack = "Acknowledgements"
    lbl_ref = "References (APA Style)"
    
    btn_download = "📄 Download Document"
    info_blind_active = "ℹ️ **Blind Review Mode Active**: Author details below will be hidden in the generated Word document."
    ack_blind_note = "🔒 *Note: Acknowledgements section will be hidden/anonymized in Blind Review mode.*"

else:
    # Bahasa Indonesia UI Texts
    ui_title = "📄 JIP 2026 Article Auto-Generator"
    ui_desc = "Isi formulir di bawah ini untuk menghasilkan naskah `.docx` siap kirim sesuai aturan Jurnal Ilmiah PLATAX."
    
    ui_mode_header = "Pengaturan Mode Naskah"
    ui_mode_label = "Pilih Versi Output Document:"
    ui_mode_opts = ["Blind Review", "Final (Lengkap)"]
    ui_mode_help = "Pilih 'Blind Review' untuk submit awal, atau 'Final (Lengkap)' jika naskah sudah diterima."
    
    ui_guide_title = "📌 **Panduan Kebijakan Submission:**\n\n"
    ui_guide_body = (
        "• **Blind Review**: Wajib digunakan saat **Pengajuan Awal (Submit & Proses Review)**. Identitas penulis dan ucapan terima kasih disembunyikan otomatis.\n\n"
        "• **Final (Lengkap)**: Digunakan setelah naskah **Diterima (Accepted / Layout Final)** untuk penerbitan."
    )
    
    tab_names = ["📌 Identitas & Judul", "📝 Abstract & Abstrak", "📄 Bab 1–4", "🤝 Pernyataan & Pustaka"]
    
    lbl_title_id = "Judul Bahasa Indonesia (Maks 20 Kata)"
    lbl_title_en = "Title in English (Max 20 Words)"
    lbl_running = "Running Title (Maks 60 Karakter)"
    lbl_author_sec = "Informasi Penulis"
    lbl_authors = "Nama Penulis (Pisahkan dengan koma)"
    lbl_affil = "Afiliasi / Instansi Penulis"
    lbl_email = "Email Korespondensi"
    
    lbl_abs_id = "Abstrak (Bahasa Indonesia)"
    lbl_kw_id = "Kata Kunci (Pisahkan dengan koma)"
    lbl_abs_en = "Abstract (English)"
    lbl_kw_en = "Keywords (Pisahkan dengan koma)"
    
    lbl_ch1 = "Bab 1: Pendahuluan"
    lbl_ch2 = "Bab 2: Metode Penelitian"
    lbl_ch3 = "Bab 3: Hasil dan Pembahasan"
    lbl_ch4 = "Bab 4: Kesimpulan dan Saran"
    
    lbl_ack = "Ucapan Terima Kasih (Acknowledgements)"
    lbl_ref = "Daftar Pustaka (APA Style)"
    
    btn_download = "📄 Download Document"
    info_blind_active = "ℹ️ **Mode Blind Review Aktif**: Data penulis di bawah ini tetap dapat diisi, namun **tidak akan dimunculkan** dalam file Word hasil unduhan."
    ack_blind_note = "🔒 *Catatan: Bagian Ucapan Terima Kasih akan otomatis disembunyikan/dianonimkan pada versi Blind Review.*"

# Render Header Utama
st.title(ui_title)
st.write(ui_desc)

# Render Mode Naskah di Sidebar
st.sidebar.subheader(ui_mode_header)
versi_naskah = st.sidebar.radio(
    ui_mode_label,
    ui_mode_opts,
    index=0,
    help=ui_mode_help
)

st.sidebar.info(ui_guide_title + ui_guide_body)

# Normalisasi status mode untuk fungsi python
is_blind = "Blind" in versi_naskah

# ---------------------------------------------------------
# TAB FORMULIR INPUT
# ---------------------------------------------------------
tab1, tab2, tab3, tab4 = st.tabs(tab_names)

with tab1:
    judul_id = st.text_input(lbl_title_id, value="Keanekaragaman Ikan Gobi di Muara Sungai Tondano, Sulawesi Utara")
    judul_en = st.text_input(lbl_title_en, value="Diversity of Goby Fish in Tondano River Estuary, North Sulawesi")
    running_title = st.text_input(lbl_running, value="Keanekaragaman Ikan Gobi di Muara Sungai Tondano")
    
    st.markdown("---")
    st.subheader(lbl_author_sec)
    
    if is_blind:
        st.info(info_blind_active)
        
    nama_penulis = st.text_area(lbl_authors, value="Ari Berty Rondonuwu*, John Doe, Jane Smith")
    afiliasi = st.text_area(lbl_affil, value="Fakultas Perikanan dan Ilmu Kelautan, Universitas Sam Ratulangi, Manado")
    email_korespondensi = st.text_input(lbl_email, value="aribertyrondonuwu@unsrat.ac.id")

with tab2:
    abstrak_id = st.text_area(lbl_abs_id, height=150, value="Penelitian ini bertujuan untuk menganalisis keanekaragaman ikan Gobi di muara Sungai Tondano...")
    kata_kunci_id = st.text_input(lbl_kw_id, value="Ikan Gobi, Keanekaragaman, Sungai Tondano, Estuari")
    abstract_en = st.text_area(lbl_abs_en, height=150, value="This study aims to analyze the diversity of Goby fish in the Tondano River estuary...")
    keywords_en = st.text_input(lbl_kw_en, value="Goby fish, Diversity, Tondano River, Estuary")

with tab3:
    bab1 = st.text_area(lbl_ch1, height=150, value="Muara Sungai Tondano memiliki peranan penting secara ekologis...")
    bab2 = st.text_area(lbl_ch2, height=150, value="Penelitian dilaksanakan pada bulan Januari hingga Maret 2026...")
    bab3 = st.text_area(lbl_ch3, height=150, value="Berdasarkan hasil tangkapan di tiga stasiun pengamatan, ditemukan...")
    bab4 = st.text_area(lbl_ch4, height=150, value="Keanekaragaman ikan Gobi di muara Sungai Tondano tergolong...")

with tab4:
    ucapan_terima_kasih = st.text_area(lbl_ack, height=100, value="Penulis mengucapkan terima kasih kepada Laboratorium Biologi Laut...")
    if is_blind:
        st.caption(ack_blind_note)
        
    daftar_pustaka = st.text_area(lbl_ref, height=200, value="Rondonuwu, A. B. (2026). Ikhtiofauna Perairan Tawar dan Payau. Jurnal Ilmiah PLATAX, 14(1), 10-20.")

# ---------------------------------------------------------
# FUNGSI PENYUSUN DOKUMEN WORD (.DOCX)
# ---------------------------------------------------------
def create_docx():
    doc = Document()
    
    # Judul Bahasa Indonesia
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run(judul_id.upper())
    run_title.bold = True
    run_title.font.size = Pt(14)
    
    # Judul Bahasa Inggris
    p_title_en = doc.add_paragraph()
    p_title_en.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title_en = p_title_en.add_run(f"({judul_en})")
    run_title_en.italic = True
    run_title_en.font.size = Pt(11)
    
    # Running Title
    p_rt = doc.add_paragraph()
    p_rt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_rt = p_rt.add_run(f"Running Title: {running_title}")
    r_rt.font.size = Pt(9)
    r_rt.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_paragraph()
    
    # LOGIKA BLIND REVIEW VS FINAL
    if not is_blind:
        p_auth = doc.add_paragraph()
        p_auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_auth = p_auth.add_run(nama_penulis)
        r_auth.bold = True
        
        p_aff = doc.add_paragraph()
        p_aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_aff = p_aff.add_run(f"{afiliasi}\nEmail: {email_korespondensi}")
        r_aff.font.size = Pt(9)
        r_aff.font.italic = True
    else:
        p_anon = doc.add_paragraph()
        p_anon.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_anon = p_anon.add_run("[INFORMASI PENULIS DIKOSONGKAN UNTUK PROSES BLIND REVIEW]")
        r_anon.font.italic = True
        r_anon.font.color.rgb = RGBColor(180, 0, 0)
        
    doc.add_paragraph()
    
    # Abstrak Bahasa Indonesia
    doc.add_heading("ABSTRAK", level=2)
    doc.add_paragraph(abstrak_id)
    p_kw = doc.add_paragraph()
    r_kw_title = p_kw.add_run("Kata Kunci: ")
    r_kw_title.bold = True
    p_kw.add_run(kata_kunci_id)
    
    doc.add_paragraph()
    
    # Abstract Bahasa Inggris
    doc.add_heading("ABSTRACT", level=2)
    doc.add_paragraph(abstract_en)
    p_kw_en = doc.add_paragraph()
    r_kw_title_en = p_kw_en.add_run("Keywords: ")
    r_kw_title_en.bold = True
    p_kw_en.add_run(keywords_en)
    
    doc.add_paragraph()
    
    # Isi Bab Utama
    doc.add_heading("1. PENDAHULUAN / INTRODUCTION", level=1)
    doc.add_paragraph(bab1)
    
    doc.add_heading("2. METODE PENELITIAN / RESEARCH METHODS", level=1)
    doc.add_paragraph(bab2)
    
    doc.add_heading("3. HASIL DAN PEMBAHASAN / RESULTS AND DISCUSSION", level=1)
    doc.add_paragraph(bab3)
    
    doc.add_heading("4. KESIMPULAN DAN SARAN / CONCLUSION AND SUGGESTIONS", level=1)
    doc.add_paragraph(bab4)
    
    # Ucapan Terima Kasih
    doc.add_heading("UCAPAN TERIMA KASIH / ACKNOWLEDGEMENTS", level=1)
    if not is_blind:
        doc.add_paragraph(ucapan_terima_kasih)
    else:
        p_thx = doc.add_paragraph("[Ucapan terima kasih disembunyikan untuk proses Blind Review]")
        p_thx.runs[0].font.italic = True
        
    # Daftar Pustaka
    doc.add_heading("DAFTAR PUSTAKA / REFERENCES", level=1)
    doc.add_paragraph(daftar_pustaka)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ---------------------------------------------------------
# TOMBOL GENERATE & DOWNLOAD
# ---------------------------------------------------------
st.markdown("---")
st.subheader(f"📥 Unduh Naskah / Download Manuscript ({versi_naskah})")

filename = "Naskah_Blind_Review.docx" if is_blind else "Naskah_Final_Lengkap.docx"

st.download_button(
    label=f"{btn_download} ({versi_naskah})",
    data=create_docx(),
    file_name=filename,
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    use_container_width=True
)
