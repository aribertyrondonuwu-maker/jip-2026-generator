"""
platax_builder.py — Versi Lengkap & Diperbaiki
Generator dokumen Word untuk Jurnal Ilmiah PLATAX 2026
"""
import io
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree

try:
    from latex2mathml.converter import convert as latex_to_mathml
    HAS_LATEX = True
except:
    HAS_LATEX = False

# =============================================================================
# KONSTANTA
# =============================================================================
FONT_NAME = 'Cambria'
BLUE = RGBColor(0x00, 0x7B, 0xB8)
BLACK = RGBColor(0x00, 0x00, 0x00)
LINK_BLUE = RGBColor(0x21, 0x96, 0xD1)


# =============================================================================
# CLASS PENULIS & NASKAH
# =============================================================================
class Penulis:
    def __init__(self, nama="", afiliasi_ids=None, is_corresp=False, email="", orcid="", **kw):
        self.nama = nama
        self.afiliasi_ids = afiliasi_ids or []
        self.is_corresp = is_corresp
        self.email = email
        self.orcid = orcid


class Naskah:
    def __init__(self, **kwargs):
        for k, v in kwargs.items():
            setattr(self, k, v)
        self.penulis_list = kwargs.get('penulis_list', [])
        self.afiliasi_list = kwargs.get('afiliasi_list', [])
        self.tabel_list = kwargs.get('tabel_list', [])
        self.gambar_list = kwargs.get('gambar_list', [])
        self.equations = {k: v for k, v in kwargs.items() if k.startswith('eq_')}


# =============================================================================
# FUNGSI XML UNTUK SET FONT & FORMATTING (PALING PENTING!)
# =============================================================================
def set_run_xml(run, size=None, bold=None, italic=None, color=None, name=FONT_NAME):
    """
    Set font properties via XML - cara paling reliable untuk python-docx.
    """
    rPr = run._r.get_or_add_rPr()
    
    # 1. Font name (Cambria)
    if name:
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:ascii'), name)
        rFonts.set(qn('w:hAnsi'), name)
        rFonts.set(qn('w:eastAsia'), name)
        rFonts.set(qn('w:cs'), name)
    
    # 2. Font size (dalam half-points, jadi 9pt = 18)
    if size is not None:
        sz_val = str(int(size * 2))
        
        # w:sz untuk Latin fonts
        sz = rPr.find(qn('w:sz'))
        if sz is None:
            sz = parse_xml(f'<w:sz {nsdecls("w")}/>')
            rPr.append(sz)
        sz.set(qn('w:val'), sz_val)
        
        # w:szCs untuk Complex Script fonts
        szCs = rPr.find(qn('w:szCs'))
        if szCs is None:
            szCs = parse_xml(f'<w:szCs {nsdecls("w")}/>')
            rPr.append(szCs)
        szCs.set(qn('w:val'), sz_val)
    
    # 3. Bold
    if bold is not None:
        b_elem = rPr.find(qn('w:b'))
        if bold:
            if b_elem is None:
                b_elem = parse_xml(f'<w:b {nsdecls("w")}/>')
                rPr.append(b_elem)
        else:
            if b_elem is not None:
                rPr.remove(b_elem)
    
    # 4. Italic
    if italic is not None:
        i_elem = rPr.find(qn('w:i'))
        if italic:
            if i_elem is None:
                i_elem = parse_xml(f'<w:i {nsdecls("w")}/>')
                rPr.append(i_elem)
        else:
            if i_elem is not None:
                rPr.remove(i_elem)
    
    # 5. Color (RGB)
    if color is not None:
        color_elem = rPr.find(qn('w:color'))
        if color_elem is None:
            color_elem = parse_xml(f'<w:color {nsdecls("w")}/>')
            rPr.append(color_elem)
        color_elem.set(qn('w:val'), f'{color[0]:02X}{color[1]:02X}{color[2]:02X}')


# =============================================================================
# FUNGSI BANTU PARAGRAF
# =============================================================================
def clear_paragraph_runs(para):
    """Hapus semua run kecuali yang pertama, lalu kosongkan."""
    runs = list(para.runs)
    if len(runs) > 1:
        for run in runs[1:]:
            run._r.getparent().remove(run._r)
    if runs:
        runs[0].text = ""
    return para


def add_formatted_text(para, text, size=9, bold=False, italic=False, 
                       color=BLACK, alignment=None, name=FONT_NAME):
    """Tambahkan teks dengan formatting lengkap ke paragraf."""
    if alignment is not None:
        para.alignment = alignment
    
    run = para.add_run(text)
    set_run_xml(run, size=size, bold=bold, italic=italic, color=color, name=name)
    return run


def fill_paragraph(para, text, size=9, bold=False, italic=False, 
                   color=BLACK, alignment=None, name=FONT_NAME):
    """
    Isi ulang paragraf dengan teks baru dan formatting lengkap.
    Ini adalah fungsi utama untuk mengganti placeholder di template.
    """
    clear_paragraph_runs(para)
    
    if alignment is not None:
        para.alignment = alignment
    
    if para.runs:
        run = para.runs[0]
    else:
        run = para.add_run("")
    
    run.text = text
    set_run_xml(run, size=size, bold=bold, italic=italic, color=color, name=name)
    return run


def find_paragraph(doc, keywords):
    """Cari paragraf yang mengandung salah satu keyword."""
    for para in doc.paragraphs:
        for kw in keywords:
            if kw in para.text:
                return para
    return None


# =============================================================================
# FUNGSI TABEL DENGAN BORDER HORIZONTAL SAJA (OPEN TABLE)
# =============================================================================
def set_table_open_borders(table):
    """
    Set tabel dengan border hanya horizontal (open table style).
    Sesuai template PLATAX: garis atas, bawah, dan antar baris.
    """
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    
    # Definisi border XML
    borders_xml = (
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    borders = parse_xml(borders_xml)
    
    # Hapus border lama jika ada
    old_borders = tblPr.find(qn('w:tblBorders'))
    if old_borders is not None:
        tblPr.remove(old_borders)
    
    tblPr.append(borders)


def create_table_in_doc(doc, headers, rows, caption_id, caption_en, footnote=""):
    """
    Buat tabel baru di dokumen dengan caption dan border yang benar.
    """
    # 1. Caption Bahasa Indonesia
    cap_para = doc.add_paragraph()
    cap_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_cap = cap_para.add_run(f"Tabel 1. {caption_id}")
    set_run_xml(run_cap, size=7.6, bold=True, name=FONT_NAME)
    
    # 2. Caption English
    cap_en_para = doc.add_paragraph()
    cap_en_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_cap_en = cap_en_para.add_run(f"Table 1. {caption_en}")
    set_run_xml(run_cap_en, size=7.6, bold=True, italic=True, name=FONT_NAME)
    
    # 3. Buat tabel
    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 4. Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(header)
        set_run_xml(run, size=9, bold=True, name=FONT_NAME)
    
    # 5. Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < num_cols:
                cell = table.rows[row_idx + 1].cells[col_idx]
                cell.text = ""
                run = cell.paragraphs[0].add_run(str(cell_text))
                set_run_xml(run, size=9, name=FONT_NAME)
    
    # 6. Set border (open table)
    set_table_open_borders(table)
    
    # 7. Footnote/keterangan
    if footnote:
        fn_para = doc.add_paragraph()
        run_fn = fn_para.add_run(f"Keterangan: {footnote}")
        set_run_xml(run_fn, size=8, italic=True, name=FONT_NAME)
    
    return table


# =============================================================================
# FUNGSI UTAMA: BANGUN DOKUMEN
# =============================================================================
def bangun(naskah=None, template_path=None, *args, **kwargs):
    """
    Fungsi utama untuk membangun dokumen Word dari template.
    """
    # Handle positional argument
    if naskah is None and args:
        naskah = args[0]
    
    is_blind = getattr(naskah, 'blind', False)
    is_en = getattr(naskah, 'bahasa', 'id') == 'en'
    
    # Load template
    tpl = template_path or "Template_Artikel_PLATAX_2026_FINAL_OJS.docx"
    if not Path(tpl).exists():
        raise FileNotFoundError(f"Template tidak ditemukan: {tpl}")
    
    doc = Document(tpl)
    
    # Set default font untuk seluruh dokumen
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(9)
    
    # Blind review metadata
    if is_blind:
        doc.core_properties.author = ""
        doc.core_properties.last_modified_by = ""
    
    # =========================================================================
    # 1. JUDUL
    # =========================================================================
    p = find_paragraph(doc, ["[Judul penelitian dalam bahasa Indonesia:"])
    if p:
        fill_paragraph(p, getattr(naskah, 'judul_id', '').upper(), 
                      size=15, bold=True, color=BLUE, 
                      alignment=WD_ALIGN_PARAGRAPH.LEFT, name=FONT_NAME)
    
    p = find_paragraph(doc, ["[Complete Research Title in English:"])
    if p:
        fill_paragraph(p, getattr(naskah, 'judul_en', ''), 
                      size=10, bold=True, italic=True, color=BLACK,
                      alignment=WD_ALIGN_PARAGRAPH.LEFT, name=FONT_NAME)
    
    p = find_paragraph(doc, ["[Judul singkat, maks. 60 karakter]"])
    if p:
        fill_paragraph(p, getattr(naskah, 'running_title', ''), 
                      size=9, name=FONT_NAME)
    
    # =========================================================================
    # 2. PENULIS & AFILIASI
    # =========================================================================
    if not is_blind and naskah.penulis_list:
        # Format nama penulis dengan nomor afiliasi
        parts = []
        for p_auth in naskah.penulis_list:
            nama = p_auth.nama if hasattr(p_auth, 'nama') else str(p_auth.get("nama", ""))
            aff = p_auth.afiliasi_ids if hasattr(p_auth, 'afiliasi_ids') else p_auth.get("afiliasi_ids", [1])
            corr = p_auth.is_corresp if hasattr(p_auth, 'is_corresp') else p_auth.get("is_corresp", False)
            aff_str = "".join(str(a) for a in aff) + ("*" if corr else "")
            parts.append(f"{nama}{aff_str}")
        
        p = find_paragraph(doc, ["[Nama Penulis 1]"])
        if p:
            fill_paragraph(p, ", ".join(parts), size=10, bold=True, color=BLACK, name=FONT_NAME)
        
        # Afiliasi
        aff_lines = []
        for i, a in enumerate(naskah.afiliasi_list, 1):
            aff_lines.append(f"{i} {a}")
        aff_text = "\n".join(aff_lines)
        
        p = find_paragraph(doc, ["[Departemen, Fakultas"])
        if p:
            fill_paragraph(p, aff_text, size=8, italic=True, color=BLACK, name=FONT_NAME)
    
    elif is_blind:
        p = find_paragraph(doc, ["[Nama Penulis 1]"])
        if p:
            fill_paragraph(p, "[ANONIM]", size=10, bold=True, name=FONT_NAME)
        
        p = find_paragraph(doc, ["[Departemen, Fakultas"])
        if p:
            fill_paragraph(p, "[AFILIASI DISEMBUNYIKAN]", size=8, italic=True, name=FONT_NAME)
    
    # =========================================================================
    # 3. ABSTRAK & KATA KUNCI
    # =========================================================================
    # Abstract English
    p = find_paragraph(doc, ["[Tulis abstract bahasa Inggris di sini."])
    if p:
        fill_paragraph(p, getattr(naskah, 'abstrak_en', ''), size=8, name=FONT_NAME)
    
    p = find_paragraph(doc, ["Keywords: [keyword 1"])
    if p:
        fill_paragraph(p, f"Keywords: {getattr(naskah, 'kata_kunci_en', '')}", size=8, name=FONT_NAME)
    
    # Abstrak Indonesia
    p = find_paragraph(doc, ["[Tulis abstrak bahasa Indonesia di sini"])
    if p:
        fill_paragraph(p, getattr(naskah, 'abstrak_id', ''), size=8, name=FONT_NAME)
    
    p = find_paragraph(doc, ["Kata kunci: [kata kunci 1"])
    if p:
        fill_paragraph(p, f"Kata kunci: {getattr(naskah, 'kata_kunci_id', '')}", size=8, name=FONT_NAME)
    
    # =========================================================================
    # 4. FOOTNOTE / KORESPONDENSI
    # =========================================================================
    if not is_blind:
        corresp = next((p for p in naskah.penulis_list 
                        if (hasattr(p, 'is_corresp') and p.is_corresp) or 
                           (isinstance(p, dict) and p.get("is_corresp"))), None)
        cn = corresp.nama if corresp and hasattr(corresp, 'nama') else "N/A"
        ft = (f"*Penulis korespondensi (Corresponding author): {cn}\n"
              f"{getattr(naskah, 'email_korespondensi', '')}\n"
              f"Tel: {getattr(naskah, 'telepon', '')}, E-mail: {getattr(naskah, 'email_korespondensi', '')}")
        
        p = find_paragraph(doc, ["*Penulis korespondensi (Corresponding author):"])
        if p:
            fill_paragraph(p, ft, size=7.6, name=FONT_NAME)
    else:
        p = find_paragraph(doc, ["*Penulis korespondensi (Corresponding author):"])
        if p:
            fill_paragraph(p, "*Informasi disembunyikan untuk blind review", size=7.6, name=FONT_NAME)
    
    # =========================================================================
    # 5. BAB 1-5 (KONTEN UTAMA)
    # =========================================================================
    sections = [
        ("1. Pendahuluan", getattr(naskah, 'bab1', '')),
        ("2.1. Waktu dan lokasi penelitian", getattr(naskah, 'metode_21', '')),
        ("2.2. Pengumpulan data", getattr(naskah, 'metode_22', '')),
        ("2.2.1. Analisis laboratorium", getattr(naskah, 'metode_221', '')),
        ("2.3. Analisis data", getattr(naskah, 'metode_23', '')),
        ("3. Hasil", getattr(naskah, 'bab3', '')),
        ("4. Pembahasan", getattr(naskah, 'bab4', '')),
        ("5. Simpulan", getattr(naskah, 'bab5', '')),
    ]
    
    for heading, content in sections:
        if content and content.strip():
            # Cari paragraf yang mengandung heading atau placeholder
            p = find_paragraph(doc, [heading, "[Paragraf"])
            if p:
                fill_paragraph(p, content, size=9, name=FONT_NAME)
    
    # =========================================================================
    # 6. PERNYATAAN AKHIR
    # =========================================================================
    closing = [
        ("Konflik kepentingan (Competing interests)", getattr(naskah, 'konflik', '')),
        ("Sumber dana (Funding sources)", getattr(naskah, 'dana', '')),
        ("Ucapan terima kasih (Acknowledgements)", "" if is_blind else getattr(naskah, 'ucapan', getattr(naskah, 'ucapan_terima_kasih', ''))),
        ("Kontribusi penulis (Authors' contributions)", getattr(naskah, 'kontribusi', '')),
        ("Ketersediaan data (Availability of data and materials)", getattr(naskah, 'data_avail', '')),
        ("Persetujuan etik (Ethics approval and consent to participate)", getattr(naskah, 'etik', '')),
    ]
    
    for label, content in closing:
        if content and content.strip():
            p = find_paragraph(doc, [label])
            if p:
                fill_paragraph(p, content.strip(), size=9, name=FONT_NAME)
    
    # =========================================================================
    # 7. DAFTAR PUSTAKA
    # =========================================================================
    refs = getattr(naskah, 'daftar_pustaka', '')
    if refs and refs.strip():
        p = find_paragraph(doc, ["[APA 7th Edition"])
        if p:
            fill_paragraph(p, refs.strip(), size=9, name=FONT_NAME)
    
    # =========================================================================
    # 8. TABEL (dari input user)
    # =========================================================================
    if naskah.tabel_list:
        for tbl_data in naskah.tabel_list:
            data_text = tbl_data.get('data', '')
            cap_id = tbl_data.get('cap_id', '')
            cap_en = tbl_data.get('cap_en', '')
            footnote = tbl_data.get('catatan', '')
            
            if data_text:
                # Parse data tabel (format: baris dipisah \n, kolom dipisah ;)
                lines = data_text.strip().split('\n')
                headers = [h.strip() for h in lines[0].split(';')]
                rows = []
                for line in lines[1:]:
                    if line.strip():
                        rows.append([c.strip() for c in line.split(';')])
                
                # Hapus tabel placeholder jika ada
                for tbl in doc.tables:
                    if any('[Parameter]' in cell.text for row in tbl.rows for cell in row.cells):
                        tbl._tbl.getparent().remove(tbl._tbl)
                        break
                
                # Buat tabel baru
                create_table_in_doc(doc, headers, rows, cap_id, cap_en, footnote)
    
    # =========================================================================
    # 9. SIMPAN DOKUMEN
    # =========================================================================
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output


# Alias untuk kompatibilitas
build_docx = bangun
