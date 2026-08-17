"""
JIP 2026 Article Auto-Generator — Jurnal Ilmiah PLATAX
Disesuaikan dengan Template_Artikel_PLATAX_2026_FINAL_OJS.docx
dan Petunjuk_Penggunaan_Template_PLATAX_2026_Bilingual_v2.docx
"""

import io
import re

import streamlit as st
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# =========================================================
# KONSTANTA TEMPLATE PLATAX 2026
# =========================================================
FONT = "Cambria"
BIRU = RGBColor(0x00, 0x7B, 0xB8)      # judul, DOI, superskrip afiliasi
BIRU_SITIRAN = RGBColor(0x21, 0x96, 0xD1)  # sitiran dalam naskah
HITAM = RGBColor(0x00, 0x00, 0x00)

OJS_URL = "https://ejournal.unsrat.ac.id/v3/index.php/platax"
DOI_DUMMY = "https://doi.org/10.35800/jip.vXXiX.XXXXX"

# Ukuran huruf per elemen (Petunjuk §2)
SZ_JUDUL_ID = 15
SZ_JUDUL_EN = 10
SZ_RUNNING = 9
SZ_PENULIS = 9
SZ_AFILIASI = 9
SZ_ABSTRAK = 8
SZ_HEADING = 8.5
SZ_SUBHEADING = 8
SZ_BODY = 9
SZ_KECIL = 7.6   # tabel, caption, pustaka, header, footer, catatan kaki

st.set_page_config(page_title="JIP 2026 Article Auto-Generator", page_icon="📄", layout="wide")


# =========================================================
# HELPER LOW-LEVEL (python-docx tidak menyediakan API-nya)
# =========================================================
def set_columns(section, num: int, space_cm: float = 0.43):
    """Set jumlah kolom pada sebuah section."""
    cols = section._sectPr.xpath("./w:cols")[0]
    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), str(int(space_cm * 566.93)))
    cols.set(qn("w:equalWidth"), "1")


def enable_even_odd_headers(doc):
    """Aktifkan header/footer berbeda untuk halaman ganjil dan genap."""
    settings = doc.settings.element
    if not settings.xpath("./w:evenAndOddHeaders"):
        el = OxmlElement("w:evenAndOddHeaders")
        settings.append(el)


def add_page_field(paragraph, size=SZ_KECIL, bold=True):
    """Sisipkan field PAGE (nomor halaman otomatis)."""
    run = paragraph.add_run()
    style_run(run, size=size, bold=bold)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def style_run(run, size=SZ_BODY, bold=False, italic=False, color=HITAM,
              superscript=False, font=FONT):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    run.font.superscript = superscript
    # Pastikan Cambria juga dipakai untuk East Asian & Complex Script
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(attr), font)
    return run


def set_hanging_indent(paragraph, hanging_cm=0.7):
    pf = paragraph.paragraph_format
    pf.left_indent = Cm(hanging_cm)
    pf.first_line_indent = Cm(-hanging_cm)


# =========================================================
# PEWARNAAN SITIRAN OTOMATIS (Petunjuk §3)
# =========================================================
# (Froese & Pauly, 2024; Hammer et al., 2021)  -> isi biru, kurung hitam
# Rondonuwu et al. (2025)                      -> nama & tahun biru
POLA_KURUNG = re.compile(r"\(([^()]*?\b(?:1[89]|20)\d{2}[a-z]?[^()]*?)\)")
POLA_NARATIF = re.compile(
    r"\b([A-ZÀ-Þ][\w’'\-]+(?:\s+(?:&|dan|and)\s+[A-ZÀ-Þ][\w’'\-]+)?"
    r"(?:\s+et\s+al\.)?)\s+\(((?:1[89]|20)\d{2}[a-z]?)\)"
)


def add_body_paragraph(doc, teks, indent=True, warnai_sitiran=True):
    """Paragraf isi naskah: Cambria 9 pt, justified, indentasi baris pertama 0,5 cm."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.5)

    if not warnai_sitiran:
        style_run(p.add_run(teks), size=SZ_BODY)
        return p

    # Kumpulkan seluruh rentang yang harus diwarnai
    spans = []  # (mulai, selesai, warna)
    for m in POLA_KURUNG.finditer(teks):
        spans.append((m.start(1), m.end(1), BIRU_SITIRAN))
    for m in POLA_NARATIF.finditer(teks):
        if any(a <= m.start(1) < b for a, b, _ in spans):
            continue
        spans.append((m.start(1), m.end(1), BIRU_SITIRAN))   # nama penulis
        spans.append((m.start(2), m.end(2), BIRU_SITIRAN))   # tahun
    spans.sort()

    kursor = 0
    for mulai, selesai, warna in spans:
        if mulai < kursor:
            continue
        if mulai > kursor:
            style_run(p.add_run(teks[kursor:mulai]), size=SZ_BODY)
        style_run(p.add_run(teks[mulai:selesai]), size=SZ_BODY, color=warna)
        kursor = selesai
    if kursor < len(teks):
        style_run(p.add_run(teks[kursor:]), size=SZ_BODY)
    return p


def add_paragraphs(doc, blok, **kw):
    """Pecah textarea multi-baris menjadi beberapa paragraf."""
    for baris in [b.strip() for b in blok.split("\n") if b.strip()]:
        add_body_paragraph(doc, baris, **kw)


def add_heading_num(doc, teks, level=1):
    """Judul bab (8,5 pt Bold) / sub-bab (8 pt Bold Italic)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    if level == 1:
        style_run(p.add_run(teks), size=SZ_HEADING, bold=True)
    else:
        style_run(p.add_run(teks), size=SZ_SUBHEADING, bold=True, italic=True)
    return p


def add_caption(doc, nomor, teks, nomor_en, teks_en):
    """Keterangan tabel/gambar: nomor bold 7,6 pt + terjemahan Inggris pada baris kedua."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    style_run(p.add_run(nomor + " "), size=SZ_KECIL, bold=True)
    style_run(p.add_run(teks), size=SZ_BODY)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    style_run(p.add_run(nomor_en + " "), size=SZ_KECIL, bold=True, italic=True)
    style_run(p.add_run(teks_en), size=SZ_BODY, italic=True)


def _garis(cell, posisi, sz=8, val="single"):
    """Atur satu sisi garis sel (val='nil' untuk menghilangkan garis)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    el = OxmlElement(f"w:{posisi}")
    el.set(qn("w:val"), val)
    el.set(qn("w:sz"), str(sz))
    el.set(qn("w:color"), "000000")
    borders.append(el)


def add_open_table(doc, data_teks, catatan=""):
    """Tabel terbuka: hanya garis horizontal di kepala dan kaki tabel."""
    baris = [b for b in data_teks.split("\n") if b.strip()]
    if not baris:
        return
    matriks = [re.split(r"[\t;]", b.strip()) for b in baris]
    n_kolom = max(len(r) for r in matriks)
    tabel = doc.add_table(rows=len(matriks), cols=n_kolom)
    tabel.autofit = False
    lebar_kol = Cm(7.6 / n_kolom)
    for kolom in tabel.columns:
        for sel in kolom.cells:
            sel.width = lebar_kol

    # Lebar tabel & tblGrid harus ikut diset, jika tidak Word/LibreOffice
    # memakai lebar halaman dan kolom terakhir terpotong di layout dua kolom
    tbl_pr = tabel._tbl.tblPr
    for lama in tbl_pr.findall(qn("w:tblW")):
        tbl_pr.remove(lama)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(int(7.6 * 566.93)))
    tbl_pr.append(tbl_w)
    grid = tabel._tbl.find(qn("w:tblGrid"))
    if grid is not None:
        for gc in grid.findall(qn("w:gridCol")):
            gc.set(qn("w:w"), str(int(7.6 * 566.93 / n_kolom)))
    for i, r in enumerate(matriks):
        for j in range(n_kolom):
            cell = tabel.cell(i, j)
            cell.paragraphs[0].text = ""
            style_run(cell.paragraphs[0].add_run(r[j] if j < len(r) else ""),
                      size=SZ_KECIL, bold=(i == 0))
            # Hapus seluruh garis, lalu pasang hanya garis horizontal yang perlu
            tc_pr = cell._tc.get_or_add_tcPr()
            old = tc_pr.find(qn("w:tcBorders"))
            if old is not None:
                tc_pr.remove(old)
            for sisi in ("left", "right"):
                _garis(cell, sisi, sz=0, val="nil")
            if i == 0:
                _garis(cell, "top", 12)
                _garis(cell, "bottom", 8)
            if i == len(matriks) - 1:
                _garis(cell, "bottom", 12)
    if catatan:
        p = doc.add_paragraph()
        style_run(p.add_run(catatan), size=SZ_KECIL, italic=True)


def add_statement(doc, judul, isi):
    """Pernyataan akhir: judul tidak bernomor + isi."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    style_run(p.add_run(judul), size=SZ_SUBHEADING, bold=True, italic=True)
    add_body_paragraph(doc, isi, indent=False, warnai_sitiran=False)


# =========================================================
# SIDEBAR
# =========================================================
st.sidebar.header("⚙️ Settings / Pengaturan")
language = st.sidebar.selectbox("🌐 Select Language / Pilih Bahasa:", ["Bahasa Indonesia", "English"], index=0)
st.sidebar.markdown("---")

EN = language == "English"

T = {
    "title": "📄 JIP 2026 Article Auto-Generator",
    "desc": ("Fill out the form below to generate a `.docx` manuscript that follows the "
             "**PLATAX 2026 template** (Cambria, two columns, FAS section order)."
             if EN else
             "Isi formulir di bawah ini untuk menghasilkan naskah `.docx` sesuai "
             "**Template PLATAX 2026** (Cambria, dua kolom, urutan bagian FAS)."),
    "mode_header": "Manuscript Output Mode" if EN else "Pengaturan Mode Naskah",
    "mode_label": "Select Output Document Version:" if EN else "Pilih Versi Output Document:",
    "mode_opts": ["Blind Review", "Final (Full Version)"] if EN else ["Blind Review", "Final (Lengkap)"],
    "mode_help": ("Select 'Blind Review' for initial submission." if EN else
                  "Pilih 'Blind Review' untuk submit awal, atau 'Final (Lengkap)' jika naskah sudah diterima."),
    "tabs": (["📌 Title & Authors", "📝 Abstracts", "📄 Sections 1–5",
              "📊 Tables & Figures", "🤝 Statements & References"] if EN else
             ["📌 Identitas & Judul", "📝 Abstract & Abstrak", "📄 Bab 1–5",
              "📊 Tabel & Gambar", "🤝 Pernyataan & Pustaka"]),
}

st.title(T["title"])
st.write(T["desc"])

st.sidebar.subheader(T["mode_header"])
versi_naskah = st.sidebar.radio(T["mode_label"], T["mode_opts"], index=0, help=T["mode_help"])
is_blind = "Blind" in versi_naskah

st.sidebar.info(
    "📌 **Blind Review**: nama penulis, afiliasi, surel, ucapan terima kasih, dan ORCID "
    "dikosongkan otomatis (Petunjuk §9).\n\n"
    "📌 **Final**: seluruh identitas dimunculkan — hanya dipakai setelah naskah ACCEPTED."
)

# =========================================================
# FORMULIR
# =========================================================
tab1, tab2, tab3, tab5, tab4 = st.tabs(T["tabs"])

with tab1:
    c1, c2 = st.columns(2)
    with c1:
        judul_id = st.text_input("Judul Bahasa Indonesia (maks. 20 kata)",
                                 value="Keanekaragaman Ikan Gobi di Muara Sungai Tondano, Sulawesi Utara")
    with c2:
        judul_en = st.text_input("Title in English (max. 20 words)",
                                 value="Diversity of Goby Fish in Tondano River Estuary, North Sulawesi")

    running_title = st.text_input("Judul singkat / Running title (maks. 60 karakter)",
                                  value="Keanekaragaman Ikan Gobi Muara Tondano",
                                  max_chars=60)
    st.caption(f"{len(running_title)}/60 karakter")

    st.markdown("---")
    st.subheader("Informasi Penulis" if not EN else "Author Information")
    if is_blind:
        st.info("ℹ️ Mode Blind Review aktif: data di bawah tetap boleh diisi, "
                "tetapi **tidak dimunculkan** dalam file Word. Simpan untuk Title Page terpisah.")

    nama_penulis = st.text_input("Nama penulis (pisahkan dengan koma, tanpa gelar)",
                                 value="Ari Berty Rondonuwu, John Doe, Jane Smith")
    afiliasi = st.text_area(
        "Afiliasi — satu baris per penulis (Departemen, Fakultas, Universitas, Kota, Kode Pos, Negara)",
        height=90,
        value=("Program Studi Ilmu Kelautan, Fakultas Perikanan dan Ilmu Kelautan, "
               "Universitas Sam Ratulangi, Manado, 95115, Indonesia"))
    penulis_korespondensi = st.text_input("Nama penulis korespondensi", value="Ari Berty Rondonuwu")
    telepon = st.text_input("Telepon korespondensi", value="+62-8XXXXXXXXXX")
    email_korespondensi = st.text_input("Surel korespondensi", value="nama@unsrat.ac.id")

with tab2:
    st.caption("Urutan template: **Abstract (EN) lebih dahulu**, baru Abstrak (ID). 150–250 kata per bahasa.")
    abstract_en = st.text_area("Abstract (English)", height=170,
                               value="This study analysed the diversity of goby fish in the Tondano River estuary...")
    keywords_en = st.text_input("Keywords (5, pisahkan dengan koma)",
                                value="Goby fish, Diversity, Tondano River, Estuary, North Sulawesi")
    st.caption(f"Abstract: {len(abstract_en.split())} kata")

    abstrak_id = st.text_area("Abstrak (Bahasa Indonesia)", height=170,
                              value="Penelitian ini menganalisis keanekaragaman ikan gobi di muara Sungai Tondano...")
    kata_kunci_id = st.text_input("Kata kunci (5, pisahkan dengan koma)",
                                  value="Ikan gobi, Keanekaragaman, Sungai Tondano, Estuari, Sulawesi Utara")
    st.caption(f"Abstrak: {len(abstrak_id.split())} kata")

with tab3:
    st.caption("Struktur FAS — Hasil dan Pembahasan **terpisah**. Satu baris kosong = paragraf baru. "
               "Sitiran seperti (Froese & Pauly, 2024) atau Rondonuwu et al. (2025) otomatis diwarnai biru.")

    bab1 = st.text_area("1. Pendahuluan", height=140,
                        value="Muara Sungai Tondano memiliki peranan ekologis penting (Carpenter & Niem, 1998).\n"
                              "Rondonuwu et al. (2025) mencatat kebaruan jenis di kawasan ini.")
    st.markdown("**2. Bahan dan Metode**")
    metode_21 = st.text_area("2.1. Waktu dan lokasi penelitian", height=90,
                             value="Penelitian dilaksanakan Januari–Maret 2026 pada tiga stasiun...")
    metode_22 = st.text_area("2.2. Pengumpulan data", height=90,
                             value="Pengambilan sampel menggunakan jaring insang dengan tiga ulangan...")
    metode_23 = st.text_area("2.3. Analisis data", height=90,
                             value="Indeks keanekaragaman Shannon-Wiener dihitung menggunakan PAST v4.03...")
    bab3 = st.text_area("3. Hasil", height=140,
                        value="Tercatat 12 jenis dari tiga stasiun pengamatan (Tabel 1).")
    bab4 = st.text_area("4. Pembahasan", height=140,
                        value="Tingginya keanekaragaman diduga berkaitan dengan variasi salinitas...")
    bab5 = st.text_area("5. Simpulan (maks. 150 kata)", height=100,
                        value="Keanekaragaman ikan gobi di muara Sungai Tondano tergolong sedang...")
    st.caption(f"Simpulan: {len(bab5.split())} kata")

with tab5:
    st.caption("Keterangan tabel di ATAS tabel, keterangan gambar di BAWAH gambar. "
               "Naskah berbahasa Indonesia WAJIB menyertakan terjemahan Inggris pada baris kedua "
               "(Table 1. / Figure 1.) — Focus & Scope §VII.A.")

    st.markdown("**Tabel 1** (opsional)")
    cap_tabel_id = st.text_input("Keterangan Tabel 1 (Indonesia)",
                                 value="Parameter kualitas air pada dua stasiun pengamatan.")
    cap_tabel_en = st.text_input("Table 1 caption (English)",
                                 value="Water quality parameters at two observation stations.")
    tabel_data = st.text_area("Isi tabel — pisahkan kolom dengan tab atau titik koma, baris pertama = kepala tabel",
                              height=110,
                              value="Parameter;Stasiun 1;Stasiun 2\n"
                                    "Suhu (°C);28,4 ± 0,3;29,1 ± 0,5\n"
                                    "Salinitas (‰);32,1 ± 0,4;31,8 ± 0,6")
    cat_tabel = st.text_input("Keterangan kaki tabel", value="Keterangan: nilai merupakan rerata ± simpangan baku.")

    st.markdown("---")
    st.markdown("**Gambar 1** (opsional)")
    berkas_gambar = st.file_uploader("Unggah gambar (≥300 dpi; TIFF/PNG/JPEG)",
                                     type=["png", "jpg", "jpeg", "tif", "tiff"])
    cap_gambar_id = st.text_input("Keterangan Gambar 1 (Indonesia)",
                                  value="Peta lokasi penelitian di Muara Sungai Tondano.")
    cap_gambar_en = st.text_input("Figure 1 caption (English)",
                                  value="Map of the study site at the Tondano River Estuary.")
    lebar_gambar = st.slider("Lebar gambar (cm) — maks. 7,6 cm untuk satu kolom", 4.0, 7.6, 7.6, 0.2)


with tab4:
    st.caption("Tujuh pernyataan akhir wajib, urutan FAS. Bila tidak relevan tulis "
               "'Tidak berlaku / Not applicable'.")
    konflik = st.text_area("Konflik kepentingan (Competing interests)", height=68,
                           value="Penulis menyatakan tidak ada konflik kepentingan yang relevan dengan artikel ini.")
    dana = st.text_area("Sumber dana (Funding sources)", height=68, value="Tidak berlaku / Not applicable.")
    ucapan_terima_kasih = st.text_area("Ucapan terima kasih (Acknowledgements)", height=68,
                                       value="Penulis berterima kasih kepada Laboratorium Biologi Laut FPIK UNSRAT.")
    if is_blind:
        st.caption("🔒 Bagian ini otomatis disembunyikan pada versi Blind Review.")
    kontribusi = st.text_area("Kontribusi penulis (CRediT)", height=68,
                              value="ABR: conceptualization, methodology, writing — original draft; "
                                    "JD: formal analysis, visualization; JS: supervision, writing — review & editing.")
    data_avail = st.text_area("Ketersediaan data (Availability of data and materials)", height=68,
                              value="Dataset tersedia dari penulis korespondensi berdasarkan permintaan yang wajar.")
    etik = st.text_area("Persetujuan etik (Ethics approval)", height=68, value="Tidak berlaku / Not applicable.")
    orcid = st.text_area("ORCID — satu baris per penulis: Nama<TAB/koma>URL", height=68,
                         value="Ari Berty Rondonuwu, https://orcid.org/0000-0000-0000-0000")

    daftar_pustaka = st.text_area("Daftar Pustaka (APA 7, satu entri per baris, min. 20 rujukan)", height=200,
                                  value="Rondonuwu, A. B., Kepel, R. C., & Tombokan, J. L. (2025). Diversity and "
                                        "biogeography of goby. Fisheries and Aquatic Sciences, 28(10), 667–676. "
                                        "https://doi.org/10.47853/FAS.2025.e56\n"
                                        "Nei, M., & Kumar, S. (2000). Molecular evolution and phylogenetics. "
                                        "Oxford University Press.")
    n_ref = len([r for r in daftar_pustaka.split("\n") if r.strip()])
    st.caption(f"{n_ref} rujukan" + ("" if n_ref >= 20 else " — minimal 20 rujukan disyaratkan."))


# =========================================================
# PENYUSUN DOKUMEN
# =========================================================
def create_docx() -> io.BytesIO:
    doc = Document()

    # ---- Gaya dasar: Cambria 9 pt, spasi tunggal
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(SZ_BODY)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0

    # ---- SECTION 1: satu kolom (judul s.d. abstrak)
    s1 = doc.sections[0]
    s1.page_width, s1.page_height = Cm(21.0), Cm(29.7)
    s1.top_margin, s1.bottom_margin = Cm(1.2), Cm(1.2)
    s1.left_margin, s1.right_margin = Cm(1.3), Cm(1.2)
    s1.header_distance = s1.footer_distance = Cm(0.5)
    set_columns(s1, 1)
    enable_even_odd_headers(doc)

    # Baris identitas terbitan
    p = doc.add_paragraph()
    style_run(p.add_run("Jurnal Ilmiah PLATAX [Vol] ([Tahun]) [Hal. awal]–[Hal. akhir]"),
              size=SZ_KECIL, italic=True)

    # Bar biru: label jenis artikel + e-ISSN
    bar = doc.add_table(rows=1, cols=2)
    bar.autofit = True
    kiri = bar.cell(0, 0).paragraphs[0]
    style_run(kiri.add_run("ARTIKEL PENELITIAN / RESEARCH ARTICLE"),
              size=11, bold=True, font="Arial Nova Cond")
    kanan = bar.cell(0, 1).paragraphs[0]
    kanan.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_run(kanan.add_run("e-ISSN 2302-3589"), size=SZ_ABSTRAK, italic=True)
    doc.add_paragraph()

    # Judul Indonesia — Cambria Bold 15 pt, biru 007BB8, rata kiri (BUKAN uppercase)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    style_run(p.add_run(judul_id), size=SZ_JUDUL_ID, bold=True, color=BIRU)

    # Judul Inggris — Bold Italic 10 pt hitam
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    style_run(p.add_run(judul_en), size=SZ_JUDUL_EN, bold=True, italic=True)

    # Judul singkat — label bold, isi biasa
    p = doc.add_paragraph()
    style_run(p.add_run("Judul singkat (running title): "), size=SZ_RUNNING, bold=True)
    style_run(p.add_run(running_title), size=SZ_RUNNING)
    doc.add_paragraph()

    # Penulis & afiliasi (disembunyikan pada blind review)
    daftar_afiliasi = [a.strip() for a in afiliasi.split("\n") if a.strip()]
    if not is_blind:
        p = doc.add_paragraph()
        nama_list = [n.strip() for n in nama_penulis.split(",") if n.strip()]
        for i, nama in enumerate(nama_list):
            if i:
                style_run(p.add_run(", "), size=SZ_PENULIS)
            style_run(p.add_run(nama), size=SZ_PENULIS)            # REGULAR, tanpa bold
            idx = min(i + 1, len(daftar_afiliasi)) if daftar_afiliasi else 1
            style_run(p.add_run(str(idx)), size=SZ_PENULIS, color=BIRU, superscript=True)
            if nama == penulis_korespondensi.strip():
                style_run(p.add_run("*"), size=SZ_PENULIS, color=BIRU, superscript=True)

        for i, afil in enumerate(daftar_afiliasi, start=1):
            p = doc.add_paragraph()
            style_run(p.add_run(str(i)), size=SZ_AFILIASI, italic=True, color=BIRU, superscript=True)
            style_run(p.add_run(afil), size=SZ_AFILIASI, italic=True)
    else:
        p = doc.add_paragraph()
        style_run(p.add_run("[Identitas penulis dan afiliasi dihapus untuk proses double-blind review "
                            "— lihat berkas Title Page terpisah]"),
                  size=SZ_PENULIS, italic=True)
    doc.add_paragraph()

    # Abstract (EN) lebih dahulu, lalu Abstrak (ID) — satu kolom penuh
    def blok_abstrak(label, isi, label_kw, kw):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        style_run(p.add_run(label), size=SZ_ABSTRAK, bold=True)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        style_run(p.add_run(isi.replace("\n", " ").strip()), size=SZ_ABSTRAK)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        style_run(p.add_run(label_kw), size=SZ_ABSTRAK, bold=True)
        style_run(p.add_run(kw), size=SZ_ABSTRAK)

    blok_abstrak("Abstract", abstract_en, "Keywords: ", keywords_en)
    blok_abstrak("Abstrak", abstrak_id, "Kata kunci: ", kata_kunci_id)

    # ---- SECTION 2: DUA KOLOM, dimulai di bab 1 Pendahuluan
    s2 = doc.add_section(WD_SECTION.CONTINUOUS)
    set_columns(s2, 2, space_cm=0.43)

    add_heading_num(doc, "1. Pendahuluan")
    add_paragraphs(doc, bab1)

    add_heading_num(doc, "2. Bahan dan Metode")
    add_heading_num(doc, "2.1. Waktu dan lokasi penelitian", level=2)
    add_paragraphs(doc, metode_21)
    add_heading_num(doc, "2.2. Pengumpulan data", level=2)
    add_paragraphs(doc, metode_22)
    add_heading_num(doc, "2.3. Analisis data", level=2)
    add_paragraphs(doc, metode_23)

    add_heading_num(doc, "3. Hasil")
    add_paragraphs(doc, bab3)

    # Tabel 1 — keterangan di ATAS tabel, bilingual
    if tabel_data.strip():
        add_caption(doc, "Tabel 1.", cap_tabel_id, "Table 1.", cap_tabel_en)
        add_open_table(doc, tabel_data, cat_tabel)

    # Gambar 1 — keterangan di BAWAH gambar, bilingual
    if berkas_gambar is not None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(io.BytesIO(berkas_gambar.getvalue()), width=Cm(lebar_gambar))
        add_caption(doc, "Gambar 1.", cap_gambar_id, "Figure 1.", cap_gambar_en)

    add_heading_num(doc, "4. Pembahasan")
    add_paragraphs(doc, bab4)

    add_heading_num(doc, "5. Simpulan")
    add_paragraphs(doc, bab5)

    # ---- Tujuh pernyataan akhir, urutan FAS
    add_statement(doc, "Konflik kepentingan (Competing interests)", konflik)
    add_statement(doc, "Sumber dana (Funding sources)", dana)
    if not is_blind:
        add_statement(doc, "Ucapan terima kasih (Acknowledgements)", ucapan_terima_kasih)
    else:
        add_statement(doc, "Ucapan terima kasih (Acknowledgements)",
                      "[Dihapus untuk proses double-blind review]")
    add_statement(doc, "Kontribusi penulis (Authors' contributions)", kontribusi)
    add_statement(doc, "Ketersediaan data (Availability of data and materials)", data_avail)
    add_statement(doc, "Persetujuan etik (Ethics approval and consent to participate)", etik)
    if not is_blind:
        add_statement(doc, "ORCID", "")
        for baris in [b.strip() for b in orcid.split("\n") if b.strip()]:
            p = doc.add_paragraph()
            style_run(p.add_run(baris), size=SZ_BODY)

    # ---- Daftar Pustaka: 7,6 pt, hanging indent 0,7 cm, seluruh entri hitam
    add_heading_num(doc, "Daftar Pustaka")
    for entri in [e.strip() for e in daftar_pustaka.split("\n") if e.strip()]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_hanging_indent(p, 0.7)
        style_run(p.add_run(entri), size=SZ_KECIL)

    # ---- Header halaman 2 dst. (ganjil & genap)
    penulis_pertama = (nama_penulis.split(",")[0].strip()
                       if (nama_penulis.strip() and not is_blind) else "[Penulis]")
    for section in doc.sections:
        section.different_first_page_header_footer = True

        h_odd = section.header
        h_odd.is_linked_to_previous = False
        p = h_odd.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.text = ""
        style_run(p.add_run(f"Jurnal Ilmiah PLATAX\t{running_title}"),
                  size=SZ_KECIL, bold=True, italic=True)

        h_even = section.even_page_header
        h_even.is_linked_to_previous = False
        p = h_even.paragraphs[0]
        p.text = ""
        style_run(p.add_run(f"{penulis_pertama}, et al.\tJurnal Ilmiah PLATAX"),
                  size=SZ_KECIL, bold=True, italic=True)

        # Footer ganjil: DOI kiri — OJS & nomor halaman kanan
        f_odd = section.footer
        f_odd.is_linked_to_previous = False
        p = f_odd.paragraphs[0]
        p.text = ""
        style_run(p.add_run(f"{DOI_DUMMY}\t"), size=SZ_KECIL)
        style_run(p.add_run(f"{OJS_URL}  |  "), size=SZ_KECIL, italic=True)
        add_page_field(p)

        # Footer genap: nomor halaman & OJS kiri — DOI kanan
        f_even = section.even_page_footer
        f_even.is_linked_to_previous = False
        p = f_even.paragraphs[0]
        p.text = ""
        add_page_field(p)
        style_run(p.add_run(f"  |  {OJS_URL}\t"), size=SZ_KECIL, italic=True)
        style_run(p.add_run(DOI_DUMMY), size=SZ_KECIL)

    # ---- Catatan kaki halaman 1
    # Diterapkan pada first-page footer KEDUA section: Word memakai section yang
    # mengakhiri halaman, LibreOffice memakai section pertama pada halaman itu.
    baris_footnote = [
        [("Diterima (Received): ", True), ("[DD Bulan YYYY]   ", False),
         ("Direvisi (Revised): ", True), ("[DD Bulan YYYY]   ", False),
         ("Disetujui (Accepted): ", True), ("[DD Bulan YYYY]", False)],
    ]
    if not is_blind:
        baris_footnote.append([("*Penulis korespondensi (Corresponding author): ", True),
                               (penulis_korespondensi, False)])
        if daftar_afiliasi:
            baris_footnote.append([(daftar_afiliasi[0], False)])
        baris_footnote.append([("Tel: ", True), (f"{telepon}, ", False),
                               ("E-mail: ", True), (email_korespondensi, False)])
    baris_footnote.append([(DOI_DUMMY, False)])
    baris_footnote.append([("Artikel ini merupakan artikel Open Access yang didistribusikan di bawah "
                            "ketentuan Creative Commons Attribution-NonCommercial 4.0 International "
                            "License (https://creativecommons.org/licenses/by-nc/4.0/), yang mengizinkan "
                            "penggunaan, distribusi, dan reproduksi non-komersial dalam media apa pun, "
                            "selama karya asli disitasi dengan benar.", False)])
    baris_footnote.append([("Hak Cipta © [Tahun] Penulis. Diterbitkan oleh Jurnal Ilmiah PLATAX, "
                            "Fakultas Perikanan dan Ilmu Kelautan, Universitas Sam Ratulangi.", False)])

    for section in doc.sections:
        f1 = section.first_page_footer
        f1.is_linked_to_previous = False
        f1.paragraphs[0].text = ""
        for i, runs in enumerate(baris_footnote):
            p = f1.paragraphs[0] if i == 0 else f1.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            if i == 0:
                borders = OxmlElement("w:pBdr")
                top = OxmlElement("w:top")
                top.set(qn("w:val"), "single")
                top.set(qn("w:sz"), "6")
                top.set(qn("w:color"), "000000")
                borders.append(top)
                p._p.get_or_add_pPr().append(borders)
            for teks, bold in runs:
                style_run(p.add_run(teks), size=SZ_KECIL, bold=bold)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# =========================================================
# UNDUH
# =========================================================
st.markdown("---")
st.subheader(f"📥 Unduh Naskah / Download Manuscript ({versi_naskah})")

peringatan = []
if len(running_title) > 60:
    peringatan.append("Judul singkat melebihi 60 karakter.")
if not (150 <= len(abstract_en.split()) <= 250):
    peringatan.append("Abstract (EN) di luar rentang 150–250 kata.")
if not (150 <= len(abstrak_id.split()) <= 250):
    peringatan.append("Abstrak (ID) di luar rentang 150–250 kata.")
if len(bab5.split()) > 150:
    peringatan.append("Simpulan melebihi 150 kata.")
if n_ref < 20:
    peringatan.append("Daftar pustaka kurang dari 20 rujukan.")
if peringatan:
    st.warning("Daftar periksa pra-submit:\n\n" + "\n".join(f"- {w}" for w in peringatan))

filename = "Naskah_Blind_Review_PLATAX_2026.docx" if is_blind else "Naskah_Final_PLATAX_2026.docx"
st.download_button(
    label=f"📄 Download Document ({versi_naskah})",
    data=create_docx(),
    file_name=filename,
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    use_container_width=True,
)
st.caption("Setelah dibuka di Word: tekan Ctrl+A lalu F9 untuk memperbarui nomor halaman otomatis. "
           "Pita masthead (logo & sampul edisi) tetap harus disalin dari template resmi PLATAX.")
